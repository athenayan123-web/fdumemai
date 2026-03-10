# -*- coding: utf-8 -*-
"""
更新功能变更日志 - 添加v2.2版本变更记录
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os

doc_path = r"E:\智能整理\openclaw\交付物\06_功能变更日志.docx"

# 读取现有文档
doc = Document(doc_path)

# 在文档开头（标题之后）插入新版本变更记录
# 找到合适的插入位置（在第一个表格之前）

# 先找到现有内容的位置
body = doc.element.body

# 创建新段落 - 版本标题
p_ver = doc.add_paragraph()
run = p_ver.add_run('\nv2.2 - 2026-03-10')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 126, 234)
p_ver.alignment = WD_ALIGN_PARAGRAPH.LEFT

p_type = doc.add_paragraph()
run = p_type.add_run('变更类型：重大功能升级 + 架构变更 + 品牌升级')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(150, 150, 150)

# 创建变更表格
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# 表头
headers = ['编号', '变更内容', '变更类型', '影响范围', '确认状态']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

# 变更记录
changes = [
    ['CHG-020', '首页主题变更：OpenClaw AI 全网比价智能平台，聚焦国际国内比价系统（模型API · Coding Plan · GPU）', '重大升级', '首页/全站', '已确认'],
    ['CHG-021', '新增模型API比价模块：覆盖OpenAI、Anthropic、Google、DeepSeek、Moonshot、阶跃星辰、智谱AI、百度、腾讯混元、字节火山、小米等19个模型实时定价', '功能新增', '比价中心', '已确认'],
    ['CHG-022', '新增Coding Plan比价模块：覆盖GitHub Copilot、Cursor、Windsurf、Augment Code、通义灵码、MarsCode等10个产品定价对比', '功能新增', '比价中心', '已确认'],
    ['CHG-023', '保留并强化GPU算力市场（图灵算力）：新增国际云GPU对比（AWS/GCP/Azure/Lambda），图灵算力作为国内GPU首要入口', '功能增强', 'GPU算力页', '已确认'],
    ['CHG-024', 'Python爬虫数据采集系统上线：实时采集15+厂商定价数据，输出JSON/JS格式供前端同步显示', '技术新增', '后端/数据', '已确认'],
    ['CHG-025', '一键部署模块改为Claw集合及对比：包含OpenClaw、QClaw、KimiClaw、JVSClaw、WorkBuddy、ArkClaw，含官网注册链接与部署手册', '架构变更', 'Claw集合页', '已确认'],
    ['CHG-026', '移除积分体系与阿里云相关入口', '运营调整', '全站', '已确认'],
    ['CHG-027', '访问策略变更：未注册用户可浏览基础比价信息，部分数据（国际API精确价格、历史趋势）需登录后查看，受限内容引导注册弹窗', '权限控制', '全站', '已确认'],
    ['CHG-028', '授权弹窗系统：新增登录弹窗、注册弹窗、受限内容引导弹窗三套弹窗交互', '功能新增', '弹窗系统', '已确认'],
    ['CHG-029', '表头署名变更：复旦大学2025未来信息创新学院工程管理宣委', '品牌变更', '全站顶部/底部', '已确认'],
    ['CHG-030', '原有网页保留为index.html，新比价平台为index_compare.html，支持双版本切换', '架构变更', '路由/导航', '已确认'],
    ['CHG-031', '所有厂商数据一键链接官网注册（同现有功能），比价表格每行均含"官网注册→"按钮', '功能增强', '比价表格', '已确认'],
    ['CHG-032', '品牌规范：GPU算力统一称"图灵算力"，不再出现"阿里云"字样', '品牌变更', '全站', '已确认'],
]

for row_data in changes:
    row = table.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

# 设置列宽
for row in table.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(3.5)
    row.cells[2].width = Inches(0.8)
    row.cells[3].width = Inches(1.0)
    row.cells[4].width = Inches(0.7)

# 保存
doc.save(doc_path)
print(f"功能变更日志已更新: {doc_path}")
print(f"新增 {len(changes)} 条变更记录 (CHG-020 ~ CHG-032)")
